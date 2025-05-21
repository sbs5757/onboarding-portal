
from flask import Flask, render_template_string, request
from docx import Document
from docx2pdf import convert
from datetime import datetime
import os
import smtplib
from email.message import EmailMessage
import mimetypes

app = Flask(__name__)

# === Configuration ===
TEMPLATE_PATH = "Sample_Offer_Letter_1.docx"
OUTPUT_DOCX_PATH = "Generated_Offer_Letter.docx"
OUTPUT_PDF_PATH = "Generated_Offer_Letter.pdf"
ONBOARDING_DOCS_FOLDER = "onboarding_docs"

SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
EMAIL_SENDER = "sbscorp5757@gmail.com"        # 🔁 Replace with your Gmail
EMAIL_PASSWORD = "tibp zevt vvad zrcq"         # 🔁 Replace with your App Password

# === HTML Form ===
HTML_FORM = """
<!doctype html>
<title>Candidate Onboarding</title>
<h2>Enter Candidate Details</h2>
<form method=post>
  Full Name: <input type=text name=full_name><br>
  Email ID: <input type=email name=email><br>
  Job Title: <input type=text name=job_title><br>
  Start Date: <input type=text name=start_date><br>
  Annual Salary: <input type=text name=salary><br><br>
  <input type=submit value="Generate & Send Offer Letter">
</form>
"""

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        full_name = request.form["full_name"]
        email = request.form["email"]
        job_title = request.form["job_title"]
        start_date = request.form["start_date"]
        salary = request.form["salary"]

        candidate_data = {
            "X1": datetime.today().strftime("%B %d, %Y"),
            "X2": full_name,
            "X3": job_title,
            "X4": start_date,
            "X5": salary,
            "X6": datetime.today().strftime("%B %d, %Y"),
            "X7": full_name,
            "X8": datetime.today().strftime("%B %d, %Y"),
        }

        # === Generate Offer Letter ===
        doc = Document(TEMPLATE_PATH)
        for para in doc.paragraphs:
            for key, value in candidate_data.items():
                if key in para.text:
                    para.text = para.text.replace(key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in candidate_data.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)

        doc.save(OUTPUT_DOCX_PATH)

        # === Convert DOCX to PDF ===
        try:
            convert(OUTPUT_DOCX_PATH, OUTPUT_PDF_PATH)
        except Exception as e:
            return f"<h3>Error converting to PDF: {str(e)}</h3>"

        # === Send Email with PDF Offer Letter + Onboarding Docs ===
        send_email_with_docs(email, full_name)

        return f"<h3>Offer letter (PDF) generated and sent to {email}</h3>"

    return render_template_string(HTML_FORM)

def send_email_with_docs(to_email, candidate_name):
    msg = EmailMessage()
    msg["Subject"] = f"Offer Letter and Onboarding Documents - {candidate_name}"
    msg["From"] = EMAIL_SENDER
    msg["To"] = to_email
    msg.set_content(f"Dear {candidate_name},\n\nPlease find attached your offer letter (PDF) and onboarding documents.\n\nBest regards,\nHR Team")

    # Attach the generated PDF offer letter
    attach_file_to_email(msg, OUTPUT_PDF_PATH)

    # Attach standard onboarding documents
    if os.path.isdir(ONBOARDING_DOCS_FOLDER):
        for file_name in os.listdir(ONBOARDING_DOCS_FOLDER):
            file_path = os.path.join(ONBOARDING_DOCS_FOLDER, file_name)
            attach_file_to_email(msg, file_path)

    # Send email
    with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as smtp:
        smtp.starttls()
        smtp.login(EMAIL_SENDER, EMAIL_PASSWORD)
        smtp.send_message(msg)

def attach_file_to_email(msg, file_path):
    with open(file_path, "rb") as f:
        file_data = f.read()
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type is None:
            mime_type = "application/octet-stream"
        maintype, subtype = mime_type.split("/")
        msg.add_attachment(file_data, maintype=maintype, subtype=subtype, filename=os.path.basename(file_path))

if __name__ == "__main__":
    app.run(debug=True, port=7860)
