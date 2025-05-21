
from flask import Flask, render_template_string, request
from docx import Document
from datetime import datetime
import os
import smtplib
import mimetypes
from email.message import EmailMessage
from docx2pdf import convert

app = Flask(__name__)

# === Configuration ===
TEMPLATE_PATH = "Sample_Offer_Letter_1.docx"
FILLED_DOCX = "Generated_Offer_Letter.docx"
FINAL_PDF = "Generated_Offer_Letter.pdf"
DOCS_FOLDER = "onboarding_docs"
EMAIL_SENDER = "sbscorp5757@gmail.com"
EMAIL_PASSWORD = "rqap eprl izok rcns"


# Replace this in HTML_FORM:
HTML_FORM = """
<!doctype html>
<title>Onboarding Portal</title>
<h2>Enter Candidate Details</h2>
<form method="post">
  Full Name: <input type="text" name="full_name"><br>
  Email: <input type="email" name="email"><br>
  Job Title: <input type="text" name="job_title"><br>
  Start Date: <input type="text" name="start_date"><br>
  Annual Salary: <input type="text" name="salary"><br>
  Work Location (optional): <input type="text" name="work_location" placeholder="5757 Flewellen Oaks Lane, Suite 704, Fulshear, TX, 77441"><br><br>

  <h3>Select documents to send:</h3>
  {% for doc in documents %}
    <input type="checkbox" name="selected_docs" value="{{ doc }}" checked> {{ doc }}<br>
  {% endfor %}
  <br>
  <input type="submit" value="Generate & Send">
</form>
"""





@app.route("/", methods=["GET", "POST"])
def index():
    docs = [f for f in os.listdir(DOCS_FOLDER)
            if os.path.isfile(os.path.join(DOCS_FOLDER, f))
            and f.endswith(('.pdf', '.docx')) and 'Sample_Offer_Letter' not in f]

    if request.method == "POST":
        full_name = request.form["full_name"]
        email = request.form["email"]
        job_title = request.form["job_title"]
        start_date = request.form["start_date"]
        salary = request.form["salary"]
        selected_files = request.form.getlist("selected_docs")


     # ✅ ✅ INSERT THIS LINE RIGHT HERE:
        work_location = request.form.get("work_location") or "5757 Flewellen Oaks Lane, Suite 704, Fulshear, TX, 77441"

        # ✅ And then use this updated dictionary
        replacements = {
            "X1": datetime.today().strftime("%B %d, %Y"),
            "X2": full_name,
            "X3": job_title,
            "X4": start_date,
            "X5": salary,
            "X6": datetime.today().strftime("%B %d, %Y"),
            "X7": full_name,
            "X8": datetime.today().strftime("%B %d, %Y"),
            "X10": work_location
        }


        doc = Document(TEMPLATE_PATH)
        for para in doc.paragraphs:
            for key, val in replacements.items():
                if key in para.text:
                    para.text = para.text.replace(key, val)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, val)
        doc.save(FILLED_DOCX)

        # Convert to PDF using docx2pdf
        convert(FILLED_DOCX, FINAL_PDF)

        # Email it
        send_email(email, full_name, selected_files)
        return f"<h3>Offer letter sent to {email} with selected documents.</h3>"

    return render_template_string(HTML_FORM, documents=docs)

def send_email(to_email, full_name, selected_files):
    msg = EmailMessage()
    msg['Subject'] = f"Offer Letter - {full_name}"
    msg['From'] = EMAIL_SENDER
    msg['To'] = to_email
    msg.set_content(f"Dear {full_name},\n\nPlease find attached your offer letter and onboarding documents.")

    attach_file(msg, FINAL_PDF)
    for fname in selected_files:
        fpath = os.path.join(DOCS_FOLDER, fname)
        if os.path.exists(fpath):
            attach_file(msg, fpath)

    with smtplib.SMTP('smtp.gmail.com', 587) as smtp:
        smtp.starttls()
        smtp.login(EMAIL_SENDER, EMAIL_PASSWORD)
        smtp.send_message(msg)

def attach_file(msg, file_path):
    with open(file_path, 'rb') as f:
        file_data = f.read()
        mime_type, _ = mimetypes.guess_type(file_path)
        maintype, subtype = mime_type.split('/') if mime_type else ('application', 'octet-stream')
        msg.add_attachment(file_data, maintype=maintype, subtype=subtype, filename=os.path.basename(file_path))

if __name__ == "__main__":
    app.run(debug=True, port=7860)

